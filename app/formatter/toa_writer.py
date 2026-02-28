"""Generate a properly formatted Table of Authorities as .docx.

Features:
- Dot leaders via raw XML (python-docx doesn't expose tab leader property)
- Hanging indent for multi-line entries
- Category headers (bold, underlined)
- Case names italicized
- Asterisk (*) for primary authorities
- Right-aligned page numbers
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, Twips

from app.models import Citation, CitationCategory, CourtPreset, TOAProject
from app.formatter.court_presets import PRESETS, DEFAULT_PRESET


class TOAWriter:
    """Generates a formatted Table of Authorities .docx document."""

    def __init__(self, project: TOAProject):
        self.project = project
        self.preset = project.preset or PRESETS[DEFAULT_PRESET]

    def generate(self, output_path: str | Path) -> Path:
        """Generate the TOA document.

        Args:
            output_path: Where to save the .docx file.

        Returns:
            Path to the generated file.
        """
        output_path = Path(output_path)
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.preset.font_name
        font.size = Pt(self.preset.font_size_body)

        # Create custom TOA entry style
        self._create_toa_style(doc)

        # Title
        self._add_title(doc)

        # Group citations by category
        grouped = self.project.citations_by_category()

        # Add each category section
        for cat in self.preset.categories_order:
            if cat in grouped:
                self._add_category_section(doc, cat, grouped[cat])

        doc.save(str(output_path))
        return output_path

    def _create_toa_style(self, doc: Document) -> None:
        """Create the TOAEntry paragraph style with hanging indent and dot-leader tab."""
        styles = doc.styles
        toa_style = styles.add_style("TOAEntry", 1)  # 1 = paragraph style
        toa_style.font.name = self.preset.font_name
        toa_style.font.size = Pt(self.preset.font_size_body)

        # Paragraph format: hanging indent
        pf = toa_style.paragraph_format
        pf.left_indent = Inches(self.preset.hanging_indent_inches)
        pf.first_line_indent = Inches(-self.preset.hanging_indent_inches)
        pf.space_after = Pt(0)
        pf.space_before = Pt(2)

        # Add right-aligned dot-leader tab stop via XML
        ppr = toa_style.element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            toa_style.element.append(ppr)

        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        # Position in twips: inches * 1440
        pos = int(self.preset.tab_position_inches * 1440)
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), str(pos))
        tabs.append(tab)
        ppr.append(tabs)

    def _add_title(self, doc: Document) -> None:
        """Add 'TABLE OF AUTHORITIES' title."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)
        run = para.add_run("TABLE OF AUTHORITIES")
        run.bold = True
        run.font.size = Pt(self.preset.font_size_title)
        run.font.name = self.preset.font_name

    def _add_category_section(
        self,
        doc: Document,
        category: CitationCategory,
        citations: list[Citation],
    ) -> None:
        """Add a category header and its citation entries."""
        # Category header
        header_para = doc.add_paragraph()
        header_para.space_before = Pt(12)
        header_para.space_after = Pt(6)
        run = header_para.add_run(category.value)
        run.bold = True
        run.underline = True
        run.font.size = Pt(self.preset.font_size_heading)
        run.font.name = self.preset.font_name

        # Citation entries
        for cite in citations:
            self._add_citation_entry(doc, cite)

    def _add_citation_entry(self, doc: Document, cite: Citation) -> None:
        """Add a single citation entry with dot leader and page numbers."""
        para = doc.add_paragraph(style="TOAEntry")

        # Primary authority marker
        prefix = ""
        if cite.is_primary:
            prefix = f"{self.preset.primary_marker} "

        # Build display text
        display = cite.display_name

        # For cases: italicize the case name portion
        if cite.category == CitationCategory.CASES and self.preset.case_italics:
            self._add_case_entry(para, prefix, display)
        else:
            if prefix:
                run = para.add_run(prefix)
                run.font.name = self.preset.font_name
                run.font.size = Pt(self.preset.font_size_body)
            run = para.add_run(display)
            run.font.name = self.preset.font_name
            run.font.size = Pt(self.preset.font_size_body)

        # Tab character triggers dot leader
        run = para.add_run("\t")
        run.font.name = self.preset.font_name
        run.font.size = Pt(self.preset.font_size_body)

        # Page numbers
        run = para.add_run(cite.page_display)
        run.font.name = self.preset.font_name
        run.font.size = Pt(self.preset.font_size_body)

    def _add_case_entry(self, para, prefix: str, display: str) -> None:
        """Add a case citation with the case name portion italicized.

        Splits on the volume number to separate case name from reporter citation.
        E.g., "Smith v. Jones, 123 F.3d 456 (6th Cir. 2020)"
        → "Smith v. Jones" (italic) + ", 123 F.3d 456 (6th Cir. 2020)" (normal)
        """
        import re

        # Find where the volume number starts
        vol_match = re.search(r",?\s*\d{1,4}\s+(?:[A-Z]|[FBUS])", display)
        if vol_match:
            case_name = display[: vol_match.start()].strip()
            rest = display[vol_match.start():]
        else:
            # Can't split — italicize everything before the first comma
            parts = display.split(",", 1)
            case_name = parts[0].strip()
            rest = "," + parts[1] if len(parts) > 1 else ""

        # Primary marker
        if prefix:
            run = para.add_run(prefix)
            run.font.name = self.preset.font_name
            run.font.size = Pt(self.preset.font_size_body)

        # Italicized case name
        run = para.add_run(case_name)
        run.italic = True
        run.font.name = self.preset.font_name
        run.font.size = Pt(self.preset.font_size_body)

        # Rest of citation (not italicized)
        if rest:
            run = para.add_run(rest)
            run.font.name = self.preset.font_name
            run.font.size = Pt(self.preset.font_size_body)
