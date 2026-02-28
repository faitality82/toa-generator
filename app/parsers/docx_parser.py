"""Extract text with page numbers from .docx files."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


@dataclass
class PageText:
    """Text content associated with a page number."""

    page: int
    text: str


# Approximate characters per page for fallback estimation
CHARS_PER_PAGE = 3000


def parse_docx(path: str | Path) -> list[PageText]:
    """Parse a .docx file and return text grouped by page number.

    Page detection strategy:
    1. Look for explicit <w:br w:type="page"/> breaks and <w:sectPr> section breaks
    2. Fall back to character-count estimation (~3000 chars/page)
    """
    doc = Document(str(path))
    pages: list[PageText] = []
    current_page = 1
    current_text_parts: list[str] = []

    def _flush_page():
        nonlocal current_text_parts
        text = "\n".join(current_text_parts).strip()
        if text:
            pages.append(PageText(page=current_page, text=text))
        current_text_parts = []

    # Check if document has explicit page breaks
    has_explicit_breaks = _has_page_breaks(doc)

    if has_explicit_breaks:
        # Strategy 1: Use explicit page breaks
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()

            # Check for page break in this paragraph's XML
            if _paragraph_has_page_break(paragraph):
                _flush_page()
                current_page += 1

            if para_text:
                current_text_parts.append(para_text)

            # Check for section break after paragraph
            if _paragraph_has_section_break(paragraph):
                _flush_page()
                current_page += 1

        # Flush remaining text
        _flush_page()
    else:
        # Strategy 2: Character-count estimation
        all_text = "\n".join(p.text for p in doc.paragraphs)
        if not all_text.strip():
            return []

        char_count = 0
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if not para_text:
                continue

            char_count += len(para_text)
            current_text_parts.append(para_text)

            if char_count >= CHARS_PER_PAGE:
                _flush_page()
                current_page += 1
                char_count = 0

        # Flush remaining text
        _flush_page()

    return pages


def _has_page_breaks(doc: Document) -> bool:
    """Check if the document contains any explicit page breaks."""
    body = doc.element.body
    # Look for <w:br w:type="page"/>
    breaks = body.findall(f".//{qn('w:br')}[@{qn('w:type')}='page']")
    if breaks:
        return True
    # Look for section breaks (sectPr within paragraphs, not the final one)
    sect_prs = body.findall(f".//{qn('w:pPr')}/{qn('w:sectPr')}")
    return len(sect_prs) > 0


def _paragraph_has_page_break(paragraph) -> bool:
    """Check if a paragraph contains a page break before its text content."""
    xml = paragraph._element
    # Check for <w:br w:type="page"/> within runs
    for br in xml.findall(f".//{qn('w:br')}"):
        if br.get(qn("w:type")) == "page":
            return True
    # Check for pageBreakBefore in paragraph properties
    ppr = xml.find(qn("w:pPr"))
    if ppr is not None:
        pb_before = ppr.find(qn("w:pageBreakBefore"))
        if pb_before is not None:
            val = pb_before.get(qn("w:val"), "true")
            if val.lower() != "false" and val != "0":
                return True
    return False


def _paragraph_has_section_break(paragraph) -> bool:
    """Check if a paragraph contains a section break (which implies a page break)."""
    xml = paragraph._element
    ppr = xml.find(qn("w:pPr"))
    if ppr is not None:
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is not None:
            # nextPage, oddPage, evenPage all imply page break
            sect_type = sect_pr.find(qn("w:type"))
            if sect_type is None:
                return True  # Default is nextPage
            val = sect_type.get(qn("w:val"), "nextPage")
            return val in ("nextPage", "oddPage", "evenPage")
    return False
